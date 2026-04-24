from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont

report_file = "Disk_Scheduling_Simulator_Report.docx"
image_file = "flow_diagram.png"


def create_flow_diagram(path):
    width, height = 1000, 620
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    boxes = [
        ((80, 80), (280, 160), "User Interface"),
        ((360, 80), (560, 160), "Frontend / React"),
        ((640, 80), (940, 160), "Backend / Express"),
        ((240, 260), (500, 340), "Algorithm Engine"),
        ((520, 260), (760, 340), "Metrics Calculator"),
        ((260, 430), (740, 510), "Results Visualization"),
    ]

    for (x1, y1), (x2, y2), text in boxes:
        draw.rectangle([x1, y1, x2, y2], outline="black", width=3)
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        text_x = x1 + (x2 - x1 - text_width) / 2
        text_y = y1 + (y2 - y1 - text_height) / 2
        draw.text((text_x, text_y), text, fill="black", font=font)

    arrows = [
        ((280, 120), (360, 120)),
        ((560, 120), (640, 120)),
        ((420, 160), (420, 260)),
        ((580, 160), (580, 260)),
        ((500, 340), (500, 430)),
    ]

    for start, end in arrows:
        draw.line([start, end], fill="black", width=4)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        if abs(dx) > abs(dy):
            direction = "right" if dx > 0 else "left"
        else:
            direction = "down" if dy > 0 else "up"
        arrow_size = 12
        if direction == "right":
            draw.polygon([end, (end[0] - arrow_size, end[1] - arrow_size / 2), (end[0] - arrow_size, end[1] + arrow_size / 2)], fill="black")
        elif direction == "left":
            draw.polygon([end, (end[0] + arrow_size, end[1] - arrow_size / 2), (end[0] + arrow_size, end[1] + arrow_size / 2)], fill="black")
        elif direction == "down":
            draw.polygon([end, (end[0] - arrow_size / 2, end[1] - arrow_size), (end[0] + arrow_size / 2, end[1] - arrow_size)], fill="black")
        elif direction == "up":
            draw.polygon([end, (end[0] - arrow_size / 2, end[1] + arrow_size), (end[0] + arrow_size / 2, end[1] + arrow_size)], fill="black")

    draw.text((50, 540), "Note: The arrows represent flow from user input to final rendered results.", fill="black", font=font)
    img.save(path)


def style_paragraph(para, size):
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(10)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_section(doc, title, paragraphs):
    heading = doc.add_heading(title, level=2)
    style_paragraph(heading, 14)
    heading.paragraph_format.space_after = Pt(6)

    for text in paragraphs:
        para = doc.add_paragraph(text)
        style_paragraph(para, 12)


doc = Document()

# Blank first page
blank_para = doc.add_paragraph('')
blank_para.paragraph_format.space_after = Pt(0)
doc.add_page_break()

heading = doc.add_heading('Disk Scheduling Simulator Project Report', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
style_paragraph(heading, 14)

subtitle = doc.add_paragraph('Full-Stack Educational Simulation Application')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
style_paragraph(subtitle, 14)

add_section(doc, '1. Project Overview', [
    'The Disk Scheduling Simulator is developed as an educational web application to demonstrate the behavior and performance characteristics of disk scheduling algorithms in operating systems.',
    'The simulator allows users to enter a sequence of disk access requests, specify the starting head position, and select scheduling algorithms to evaluate. The application then computes the complete seek sequence and associated metrics, enabling comparison among common disk scheduling strategies.',
    'This project is intended for operating systems learners, system designers, and anyone who wishes to understand how disk head movement affects I/O efficiency. By combining real-time visualization with metric-based analysis, the simulator makes abstract scheduling concepts tangible.',
    'The web-based architecture separates the user experience from the algorithm engine, enabling clear separation of concerns. The frontend is responsible for user interaction and visualization, while the backend executes algorithm logic and returns structured results.',
    'Disk scheduling algorithms are critical in operating systems because they determine the order in which disk I/O requests are serviced, directly impacting system performance and user experience. Poor scheduling can lead to high seek times, increased latency, and reduced throughput.',
    'The simulator supports four fundamental algorithms: First Come First Served (FCFS), Shortest Seek Time First (SSTF), SCAN, and Circular SCAN (C-SCAN). Each algorithm has unique characteristics that affect fairness, efficiency, and response times under different workload patterns.',
    'Users can experiment with various input scenarios, including different request sequences, head positions, and disk sizes, to observe how algorithm performance varies. This hands-on approach helps build intuition about algorithmic tradeoffs in real-world disk I/O management.',
    'The application provides both single-algorithm simulation and side-by-side comparison modes, making it an effective teaching tool for computer science courses and self-study. The interactive nature encourages exploration and deeper understanding of operating system concepts.',
])

doc.add_page_break()

add_section(doc, '2. Module-Wise Breakdown', [
    'This section explores the project architecture module by module, describing the responsibilities, implementation details, and how each part contributes to the overall system.',
    'A well-organized module structure simplifies maintenance and allows the project to scale with additional algorithms and visualization features over time.',
    'The modular design follows best practices for web application development, separating concerns between data processing, business logic, and user interface components.',
    'Each module is designed to be independently testable and maintainable, reducing the complexity of debugging and feature additions.',
])

add_section(doc, '2.1 Backend Architecture', [
    'The backend is implemented using Node.js and the Express framework. Its primary responsibility is to accept simulation requests from the frontend, validate input, select the appropriate disk scheduling algorithm, compute metrics, and return the results as JSON.',
    'The backend entry point is server.js. It configures Express middleware for cross-origin resource sharing (CORS) and JSON parsing, and it mounts the API router under the /api path. The backend also includes a lightweight health-check route to confirm the server is running.',
    'Server.js acts as the central configuration point, ensuring that all incoming requests are properly handled and that the application can communicate securely with the frontend during development.',
    'The backend architecture is designed to be stateless, meaning each request is processed independently without maintaining session state, which simplifies scaling and deployment.',
])

backend_components = [
    'routes/scheduling.js: Defines REST endpoints for simulation and comparison requests, then delegates processing to the controller layer. This module handles HTTP routing and ensures proper request/response formats.',
    'controllers/schedulingController.js: Validates incoming parameters, executes the requested algorithm (or all algorithms for comparison mode), and formats the result payload for the frontend. It acts as the business logic coordinator.',
    'algorithms/fcfs.js: Implements the First Come First Served algorithm, which processes requests in the order received. This simple algorithm serves as a baseline for comparison with more sophisticated methods.',
    'algorithms/sstf.js: Implements the Shortest Seek Time First algorithm, selecting the request closest to the current head position at each step. This greedy approach minimizes immediate seek time but can lead to starvation.',
    'algorithms/scan.js: Implements the SCAN algorithm, sweeping in one direction until the end of the disk and then reversing direction. This elevator-like algorithm provides more predictable service patterns.',
    'algorithms/cscan.js: Implements the C-SCAN algorithm, which moves in one direction and wraps back to the start without servicing requests on the return sweep. This variant reduces response time variance.',
    'utils/metrics.js: Calculates detailed performance metrics including total seek time, average seek time, maximum seek time, variance, standard deviation, throughput, and request count. These metrics provide quantitative evaluation of algorithm efficiency.',
]
for component in backend_components:
    para = doc.add_paragraph(component, style='List Bullet')
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

doc.add_page_break()

add_section(doc, '2.2 Frontend Architecture', [
    'The frontend is constructed with React and Vite, and it uses Tailwind CSS for design consistency. The primary objective of the frontend is to collect user input, present simulation controls, render performance results, and visualize disk head movement.',
    'The application supports both single-algorithm simulation and simultaneous comparison of multiple algorithms, making it useful for learners who want to observe differences directly.',
    'React enables the creation of reusable UI components, which improves code maintainability and allows for consistent user experiences across different parts of the application.',
    'Vite provides fast development builds and hot module replacement, significantly improving developer productivity during the coding phase.',
])

frontend_components = [
    'src/App.jsx: The root component manages application state, loading status, mode selection, and error handling. It connects user interactions from the configuration panel to the results display. This component orchestrates the overall application flow.',
    'src/components/InputPanel.jsx: Provides form fields for disk request input, maximum track number, initial head position, algorithm selection, and direction control for SCAN. It also supplies random request generation for quick experimentation. This component handles all user input validation on the client side.',
    'src/components/ResultsPanel.jsx: Displays results in either single-algorithm or comparison view. It renders summaries, charts, and tables to make algorithm performance immediately understandable. This component adapts its layout based on the selected mode.',
    'src/components/DiskChart.jsx: Builds an interactive chart that plots the seek sequence and highlights head movement through the requested tracks. The chart uses Recharts library for smooth animations and responsive design.',
    'src/components/MetricsCard.jsx: Shows individual performance measurements in a compact card format, improving readability and focus. Each card displays a specific metric with appropriate formatting and units.',
    'src/services/api.js: Encapsulates network logic for POST /api/simulate and POST /api/compare, enabling the frontend to interact with the backend cleanly and asynchronously. This module handles HTTP requests and error responses.',
]
for component in frontend_components:
    para = doc.add_paragraph(component, style='List Bullet')
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

doc.add_page_break()

add_section(doc, '3. Functionalities', [
    'This section describes the functional capabilities of the simulator in depth, covering the core user flows, data validation, algorithm execution, comparison mode, and metric reporting.',
    'The functionalities are designed to provide a comprehensive learning experience, allowing users to explore disk scheduling concepts through interactive experimentation.',
    'Each functionality is implemented with user experience in mind, providing clear feedback and intuitive controls to facilitate learning.',
    'The application handles edge cases gracefully, such as empty request lists or invalid head positions, ensuring robust operation under various input conditions.',
])

functionality_items = [
    'User Input Interface: The simulator provides a responsive form for entering an array of disk track requests, which can be typed manually or generated automatically for experimentation. The interface includes helpful placeholders and validation messages.',
    'Validation Layer: The frontend and backend work together to validate inputs. The backend checks request values, head position bounds, algorithm selection, and track range before processing. This dual validation ensures data integrity.',
    'Single Algorithm Simulation: Users can simulate FCFS, SSTF, SCAN, or C-SCAN individually. Each simulation executes a complete seek sequence and returns the path taken by the disk head. Results include both visual and numerical representations.',
    'Comparison Mode: The compare feature runs all supported algorithms against the same input set and reports results side-by-side, highlighting the algorithm with the best total seek time. This mode facilitates direct algorithmic comparison.',
    'Performance Metrics: The backend computes metrics that quantify algorithm efficiency and consistency. These metrics help compare strategies beyond raw sequence order. Metrics include statistical measures like variance and standard deviation.',
    'Visualization: The frontend displays both a chart representation and a detailed list of the seek sequence, allowing users to understand the physical movement of the disk head. The chart updates dynamically based on simulation results.',
    'Error Feedback: Invalid input scenarios are communicated clearly to the user, reducing confusion and improving usability during experimentation. Error messages are contextual and suggest corrective actions.',
]
for item in functionality_items:
    para = doc.add_paragraph(item, style='List Number')
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

doc.add_page_break()

add_section(doc, '3.1 Single Algorithm Processing Workflow', [
    'When a user chooses a single algorithm and submits the form, the frontend packages the request data into a JSON payload. The backend receives this payload and validates each field. After validation, the specified algorithm module runs and returns the computed sequence of tracks and seek distances. Finally, the backend assembles the response with the computed metrics and sends it back to the frontend for rendering.',
    'The simulation flow consists of the following steps:',
    'Input normalization and parsing of commas-separated track requests into numeric values. This step ensures consistent data format regardless of user input style.',
    'Validation to ensure every track is within bounds and the head position is valid. Invalid inputs are rejected with specific error messages.',
    'Algorithm selection logic that dispatches to the appropriate scheduling module. Each algorithm has its own implementation file for modularity.',
    'Sequence computation that records the exact order in which requests are serviced. This sequence forms the basis for metric calculations.',
    'Metric calculation based on the distances between consecutive head positions. Multiple metrics are computed to provide comprehensive performance analysis.',
    'Result packaging and JSON response delivery back to the client. The response includes all necessary data for frontend rendering.',
    'This workflow ensures that each simulation is processed efficiently and accurately, providing users with reliable results for educational purposes.',
])

doc.add_page_break()

add_section(doc, '3.2 Comparison Mode Workflow', [
    'The comparison mode runs FCFS, SSTF, SCAN, and C-SCAN on the same request data set using identical input parameters. The backend computes the result for each algorithm separately, and the frontend presents the results in a comparative table that highlights the best performing algorithm in terms of total seek time.',
    'This mode is especially valuable for educational use because it reveals the tradeoffs between fairness, throughput, and seek cost among different scheduling strategies.',
    'The comparison workflow involves parallel execution of all algorithms, ensuring that each receives the same input conditions for fair comparison.',
    'Results are aggregated and presented in a unified format, allowing users to quickly identify performance differences and algorithmic strengths.',
    'The best-performing algorithm is automatically highlighted, but users can examine all metrics to understand the full performance profile of each strategy.',
])

doc.add_page_break()

add_section(doc, '4. Technology Used', [
    'This project uses modern web development technologies to provide a smooth interactive experience while keeping the backend logic lightweight and extensible.',
    'The technology stack was chosen for its maturity, community support, and suitability for educational applications.',
    'Each technology component was selected based on its ability to support the project requirements while maintaining development efficiency.',
])

heading_41 = doc.add_heading('4.1 Programming Languages', level=3)
style_paragraph(heading_41, 14)

for language in [
    'JavaScript: Used for both the frontend and the backend. The frontend is built with React and Vite, while the backend runs on Node.js. JavaScript provides a unified language across the full stack.',
    'HTML: Provides the document structure for the web interface and serves the base page loaded by the browser. HTML5 features are used for semantic markup and accessibility.',
    'CSS: Used for layout, responsive design, and styling of the user interface, including Tailwind CSS utility classes. CSS ensures the application looks professional and works on various screen sizes.',
]:
    para = doc.add_paragraph(language)
    style_paragraph(para, 12)

doc.add_heading('4.2 Libraries and Tools', level=3)
heading_42 = doc.paragraphs[-1]
style_paragraph(heading_42, 14)

for tool in [
    'React: Enables component-based UI development and declarative rendering. React simplifies building interactive user interfaces with reusable components.',
    'Vite: Provides fast development startup, module hot replacement, and efficient build tooling. Vite improves the development experience with quick feedback loops.',
    'Tailwind CSS: Supplies low-level CSS utility classes for rapid UI construction and responsive design. Tailwind allows for consistent styling without custom CSS files.',
    'Recharts: Used for plotting the seek sequence chart and visualizing algorithm behavior. Recharts provides customizable and responsive chart components.',
    'Express: Serves as the backend framework for routing and request handling. Express simplifies building RESTful APIs with middleware support.',
    'cors: Allows the frontend to communicate with the backend API during development without cross-origin restrictions. CORS is essential for local development setups.',
]:
    para = doc.add_paragraph(tool)
    style_paragraph(para, 12)

doc.add_heading('4.3 Other Tools', level=3)
heading_43 = doc.paragraphs[-1]
style_paragraph(heading_43, 14)

for tool in [
    'npm: Package manager used to install dependencies and run development scripts for both frontend and backend. npm manages the project dependencies and scripts.',
    'GitHub: Recommended platform for version control, issue tracking, and repository management. GitHub provides collaboration features and project visibility.',
    'Visual Studio Code: Suggested development environment for editing project code, debugging, and running scripts. VS Code offers excellent JavaScript and React support.',
]:
    para = doc.add_paragraph(tool)
    style_paragraph(para, 12)

doc.add_page_break()

add_section(doc, '5. Flow Diagram', [
    'The flow diagram below illustrates the interaction among all major components in the system. It captures the lifecycle of a request from the user to the final rendered result.',
    'The diagram shows how data flows from the user interface through the frontend components, across the network to the backend, through algorithm processing, and back to the visualization layer.',
    'Each box represents a major system component, and the arrows indicate the direction of data flow and control.',
    'The flow diagram helps developers and users understand the system architecture and how different parts interact to produce the final simulation results.',
])
create_flow_diagram(image_file)
doc.add_picture(image_file, width=Inches(6.5))

doc.add_page_break()

add_section(doc, '6. Revision Tracking on GitHub', [
    'A disciplined revision tracking approach improves project transparency and makes collaboration easier. In a typical workflow, each feature or fix is developed on a separate branch, and pull requests are used to review changes before merging.',
    'A repository named disk-scheduling-simulator should contain distinct directories for frontend and backend code, a README file with instructions, and any generated reports or documentation.',
    'Useful GitHub practices for this project include meaningful commit messages, version tagging for major releases, and issue tracking for feature requests or bug reports.',
    'GitHub Actions can be used for continuous integration, automatically running tests and builds on each push to ensure code quality.',
    'The repository structure should include clear documentation, contribution guidelines, and licensing information to facilitate community involvement.',
])

doc.add_page_break()

add_section(doc, '7. Conclusion and Future Scope', [
    'The Disk Scheduling Simulator is an effective learning tool for visualizing operating system disk scheduling algorithms. It bridges theoretical concepts with practical experimentation by exposing algorithm internals through simulation and performance metrics.',
    'As an educational application, it enables learners to develop intuition about algorithmic tradeoffs, such as why SSTF may reduce seek cost but risk starvation, or why SCAN can provide more balanced service across the disk surface.',
    'Future scope includes adding more advanced scheduling policies such as LOOK, C-LOOK, and multi-level queue scheduling, as well as supporting time-stamped requests and interactive animations of head movement. Another possible future enhancement is a file upload mechanism allowing real-world disk traces to be simulated directly from actual workload data.',
    'The project could also benefit from user accounts for saving simulation configurations, a gallery of example scenarios, and integration with educational platforms for automated assessment.',
    'Overall, the simulator successfully demonstrates the importance of disk scheduling in operating systems and provides a foundation for further exploration of I/O management techniques.',
])

doc.add_page_break()

add_section(doc, '8. References', [
    'React Documentation. https://react.dev/',
    'Express Documentation. https://expressjs.com/',
    'Vite Documentation. https://vitejs.dev/',
    'Tailwind CSS Documentation. https://tailwindcss.com/',
    'Operating Systems concepts and disk scheduling theory.',
    'General web development resources for Node.js and React.',
    'Academic papers on disk scheduling algorithms and performance analysis.',
])

doc.add_page_break()

add_section(doc, 'Appendix A: AI-Generated Project Elaboration/Breakdown Report', [
    'The AI-generated report elaborates on the project architecture, algorithm implementations, dataset processing, and user interaction patterns. It highlights the separation between frontend visualization and backend algorithm execution, which ensures the simulator is maintainable and easy to extend.',
    'The backend routing layer receives simulation requests and delegates them to controller modules. Each controller validates input and calls the appropriate algorithm. Metrics are then calculated and returned with the simulation sequence, enabling the frontend to render both numeric values and graphical results.',
    'The frontend uses modular React components to support user input, chart visualization, and comparative result tables. InputPanel.jsx collects and normalizes data, while ResultsPanel.jsx handles display logic for both single and comparison modes. The design encourages a clear feedback loop for learners.',
    'The architecture balances usability with educational clarity. The frontend uses a proxy-based development server to communicate with the backend API without CORS issues, while the backend remains lightweight and focused on algorithm correctness and metric reporting.',
    'The project demonstrates good software engineering practices through its modular design, clear separation of concerns, and comprehensive error handling.',
])

doc.add_page_break()

add_section(doc, 'Appendix B: Problem Statement', [
    'Disk scheduling is a core operating systems problem where the disk head must service multiple read or write requests placed across different track positions. Efficient scheduling reduces the total seek distance and improves overall throughput, while poor scheduling can cause long delays and unfair wait times.',
    'The problem addressed by this project is to provide an interactive simulator that clearly demonstrates how different disk scheduling algorithms behave with the same input workload, and how they trade off between seek distance, request ordering, and fairness.',
    'By visualizing the actual head movement and reporting performance metrics, the simulator helps learners understand why some algorithms are more appropriate under different workload conditions and how scheduling strategy impacts the end-to-end cost of servicing disk requests.',
    'The simulator fills an educational gap by providing hands-on experience with disk scheduling concepts that are often taught theoretically in operating systems courses.',
    'This tool enables students and professionals to experiment with different scenarios, observe algorithmic behavior, and develop a deeper understanding of I/O subsystem performance characteristics.',
])

doc.save(report_file)
print(f"Generated report: {report_file}")
